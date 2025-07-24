# main.py

import os
import re
import google.generativeai as genai
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
from dotenv import load_dotenv

# Struktur Bab II
document_structure = [
    {"title": "2.1 Latar Belakang dan Sejarah Perusahaan"},
    {"title": "2.2 Profil & Struktur Perusahaan", "subsections": [
        {"title": "2.2.1 Profil Perusahaan"},
        {"title": "2.2.2 Struktur Perusahaan"},
    ]},
    {"title": "2.3 Visi, Misi dan Tujuan Perusahaan", "subsections": [
        {"title": "2.3.1 Visi Perusahaan"},
        {"title": "2.3.2 Misi Perusahaan"},
        {"title": "2.3.3 Nilai Perusahaan"},
        {"title": "2.3.4 Tujuan Perusahaan"},
    ]},
    {"title": "2.4 Prinsip – Prinsip Tata Kelola Korporasi (Good Corporate Governance)"},
    {"title": "2.5 Evaluasi Pelaksanaan RJP Sebelumnya (2021–2025)", "subsections": [
        {"title": "2.5.1 Pencapaian Tujuan Perusahaan pada RJP 2021–2025"},
        {"title": "2.5.2 Pelaksanaan strategi dan kebijakan yang telah ditetapkan"},
    ]},
    {"title": "2.6 Benchmark Perusahaan Global Terminal Operator (GTO)"}
]

# Struktur Bab III
document_structure_bab3 = [
    {"title": "3.1 *Milestone* Capaian Perusahaan"},
    {"title": "3.2 Kinerja Operasional Perusahaan"},
    {"title": "3.3 Kinerja Keuangan Perusahaan"},
    {"title": "3.4 Ketercapaian *Key Performance Indicators*"}
]

document_structure_bab4 = [
    {"title": "4.1 Analisis Strength, Weakness, Opportunity & Threat (SWOT)"},
    {"title": "4.2 Matriks TOWS dan Kesimpulan Arah Pengembangan Bisnis Perusahaan"},
    {"title": "4.3 Analisis Industri dan Daya Saing (BUMN) Pada Sektor Bisnis", "subsections": [
        {"title": "4.3.1 Analisis Lingkungan Makro"},
        {"title": "4.3.2 Analisis Industri Kepelabuhanan dan Logistik"},
        {"title": "4.3.3 Analisis PESTLE"},
        {"title": "4.3.4 Analisis Daya Saing dan Posisi"},
    ]}
]

def get_google_api_key():
    try:
        import streamlit as st
        return st.secrets.get("GOOGLE_API_KEY")
    except:
        load_dotenv()
        return os.getenv("GOOGLE_API_KEY")

def fetch_section_content(company_name, section_title, temperature=0.7, model_name="models/gemini-1.5-flash-latest"):
    prompt = f"""
Kamu adalah seorang profesional senior management consultant yang ahli dalam membuat profil perusahaan.
Tulis bagian "{section_title}" untuk profil perusahaan "{company_name}".
Gunakan bahasa formal dan komprehensif.

Tambahkan daftar sumber referensi aktual dari publikasi resmi, laporan keuangan, kementerian terkait, atau sumber berita ekonomi.
Tuliskan referensi tersebut dalam bentuk daftar bernomor dan sebutkan judul publikasi serta tahun atau tautannya bila diketahui.
"""

    try:
        model = genai.GenerativeModel(model_name=model_name)
        response = model.generate_content(prompt)
        usage = response.usage_metadata if hasattr(response, "usage_metadata") else None
        tokens_in = usage.prompt_token_count if usage else 0
        tokens_out = usage.candidates_token_count if usage else 0
        return response.text, tokens_in, tokens_out
    except Exception as e:
        return f"[GAGAL MENGAMBIL KONTEN: {e}]", 0, 0

def add_table_of_contents(paragraph):
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)

def generate_bab2_docx(company_name, temperature=0.7, model_name="models/gemini-1.5-flash-latest"):
    api_key = get_google_api_key()
    if not api_key:
        raise ValueError("API Key tidak ditemukan.")
    genai.configure(api_key=api_key)

    doc = Document()
    doc.add_heading("BAB II  PENDAHULUAN", level=1)

    toc_paragraph = doc.add_paragraph()
    add_table_of_contents(toc_paragraph)
    doc.add_paragraph("(Klik kanan pada daftar isi di Microsoft Word > Update field untuk menampilkan halaman)")

    total_tokens_in = 0
    total_tokens_out = 0

    for section in document_structure:
        doc.add_heading(section["title"], level=2)
        if "subsections" in section:
            for sub in section["subsections"]:
                doc.add_heading(sub["title"], level=3)
                content, tokens_in, tokens_out = fetch_section_content(company_name, sub["title"], temperature, model_name)
                doc.add_paragraph(content.strip())
                total_tokens_in += tokens_in
                total_tokens_out += tokens_out
        else:
            content, tokens_in, tokens_out = fetch_section_content(company_name, section["title"], temperature, model_name)
            doc.add_paragraph(content.strip())
            total_tokens_in += tokens_in
            total_tokens_out += tokens_out

    safe_name = re.sub(r'[^\w\s-]', '', company_name).replace(" ", "_")
    filename = f"Bab_II_Profil_{safe_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)

    return filename, total_tokens_in, total_tokens_out

def generate_bab3_docx(company_name, temperature=0.7, model_name="models/gemini-1.5-flash-latest"):
    api_key = get_google_api_key()
    if not api_key:
        raise ValueError("API Key tidak ditemukan.")
    genai.configure(api_key=api_key)

    doc = Document()
    doc.add_heading("BAB III  KINERJA PERUSAHAAN TAHUN 2021–2024", level=1)

    toc_paragraph = doc.add_paragraph()
    add_table_of_contents(toc_paragraph)
    doc.add_paragraph("(Klik kanan pada daftar isi di Microsoft Word > Update field untuk menampilkan halaman)")

    total_tokens_in = 0
    total_tokens_out = 0

    for section in document_structure_bab3:
        doc.add_heading(section["title"], level=2)
        content, tokens_in, tokens_out = fetch_section_content(company_name, section["title"], temperature, model_name)
        doc.add_paragraph(content.strip())
        total_tokens_in += tokens_in
        total_tokens_out += tokens_out

    safe_name = re.sub(r'[^\w\s-]', '', company_name).replace(" ", "_")
    filename = f"Bab_III_Kinerja_{safe_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)

    return filename, total_tokens_in, total_tokens_out

def generate_bab4_docx(company_name, temperature=0.7, model_name="models/gemini-1.5-flash-latest"):
    api_key = get_google_api_key()
    if not api_key:
        raise ValueError("API Key tidak ditemukan.")
    genai.configure(api_key=api_key)

    doc = Document()
    doc.add_heading("BAB IV ANALISIS POSISI PERUSAHAAN", level=1)

    toc_paragraph = doc.add_paragraph()
    add_table_of_contents(toc_paragraph)
    doc.add_paragraph("(Klik kanan pada daftar isi di Microsoft Word > Update field untuk menampilkan halaman)")

    total_tokens_in = 0
    total_tokens_out = 0

    for section in document_structure_bab4:
        doc.add_heading(section["title"], level=2)
        if "subsections" in section:
            for sub in section["subsections"]:
                doc.add_heading(sub["title"], level=3)
                content, tokens_in, tokens_out = fetch_section_content(company_name, sub["title"], temperature, model_name)
                doc.add_paragraph(content.strip())
                total_tokens_in += tokens_in
                total_tokens_out += tokens_out
        else:
            content, tokens_in, tokens_out = fetch_section_content(company_name, section["title"], temperature, model_name)
            doc.add_paragraph(content.strip())
            total_tokens_in += tokens_in
            total_tokens_out += tokens_out

    safe_name = re.sub(r'[^\w\s-]', '', company_name).replace(" ", "_")
    filename = f"Bab_IV_Analisis_Posisi_{safe_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)

    return filename, total_tokens_in, total_tokens_out
