from modules.fetch_section import fetch_section_content
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
import re

def add_table_of_contents(paragraph):
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)

def generate_docx_for_structure(company_name, document_structure, doc_title, filename_prefix, temperature=0.7, model_name="models/gemini-1.5-pro", depth=1, breadth=3):
    doc = Document()
    doc.add_heading(doc_title, level=1)
    toc_paragraph = doc.add_paragraph()
    add_table_of_contents(toc_paragraph)
    doc.add_paragraph("(Klik kanan pada daftar isi di Microsoft Word > Update field untuk menampilkan halaman)")

    total_tokens_in = 0
    total_tokens_out = 0

    for section in document_structure:
        doc.add_heading(section["title"], level=2)
        if "subsections" in section:
            for sub in section["subsections"]:
                doc.add_heading(sub["title"], level=3)
                content, urls, tokens_in, tokens_out = fetch_section_content(company_name, sub["title"], temperature, model_name, depth, breadth)
                doc.add_paragraph(content)
                if urls:
                    doc.add_paragraph("Sumber Referensi:", style='Intense Quote')
                    for url in urls:
                        doc.add_paragraph(f"- {url}", style='List Bullet')
                total_tokens_in += tokens_in
                total_tokens_out += tokens_out
        else:
            content, urls, tokens_in, tokens_out = fetch_section_content(company_name, section["title"], temperature, model_name, depth, breadth)
            doc.add_paragraph(content)
            if urls:
                doc.add_paragraph("Sumber Referensi:", style='Intense Quote')
                for url in urls:
                    doc.add_paragraph(f"- {url}", style='List Bullet')
            total_tokens_in += tokens_in
            total_tokens_out += tokens_out

    safe_name = re.sub(r'[^\w\s-]', '', company_name).replace(" ", "_")
    filename = f"{filename_prefix}_{safe_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    return filename, total_tokens_in, total_tokens_out
